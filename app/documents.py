from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

_UNSAFE_FILENAME = re.compile(r"[^a-zA-Z0-9_-]+")


def _slug(value: str) -> str:
    return _UNSAFE_FILENAME.sub("-", value.strip()).strip("-").lower()[:60] or "application"


def _configure(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Heading 1"].font.name = "Arial"
    styles["Heading 1"].font.size = Pt(13)


def write_application_documents(
    job: dict[str, Any],
    content: dict[str, Any],
    *,
    candidate_name: str,
    output_root: Path,
) -> tuple[Path, Path, Path]:
    folder = output_root / f"{job['id']}-{_slug(job['employer'])}-{_slug(job['title'])}"
    folder.mkdir(parents=True, exist_ok=True)

    cv = Document()
    _configure(cv)
    name = cv.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run(candidate_name.upper())
    run.bold = True
    run.font.size = Pt(18)
    title = cv.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run(content["cv_title"]).bold = True

    cv.add_heading("Professional Profile", level=1)
    cv.add_paragraph(content["professional_summary"])
    cv.add_heading("Core Skills", level=1)
    cv.add_paragraph(" | ".join(content["core_skills"]))
    cv.add_heading("Professional Experience", level=1)
    for role in content["experience"]:
        heading = cv.add_paragraph()
        heading.add_run(f"{role['role']} — {role['employer']}").bold = True
        heading.add_run(f"\n{role['location']} | {role['dates']}")
        for bullet in role["bullets"]:
            cv.add_paragraph(bullet, style="List Bullet")
    cv.add_heading("Education", level=1)
    for education in content["education"]:
        paragraph = cv.add_paragraph()
        paragraph.add_run(education["qualification"]).bold = True
        paragraph.add_run(
            f" — {education['institution']} | {education['dates']}"
        )
    cv_path = folder / "tailored-cv.docx"
    cv.save(cv_path)

    letter = Document()
    _configure(letter)
    letter.add_paragraph(candidate_name)
    letter.add_paragraph(f"Application: {job['title']} — {job['employer']}")
    for paragraph in content["cover_letter"].split("\n"):
        if paragraph.strip():
            letter.add_paragraph(paragraph.strip())
    letter_path = folder / "cover-letter.docx"
    letter.save(letter_path)

    notes_path = folder / "review-notes.txt"
    notes = [
        "REVIEW BEFORE SUBMISSION",
        "",
        *[f"- {note}" for note in content["review_notes"]],
        "",
        f"Vacancy: {job['vacancy_url']}",
    ]
    notes_path.write_text("\n".join(notes), encoding="utf-8")
    return cv_path, letter_path, notes_path

